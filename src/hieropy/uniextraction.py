import re
from lxml import html, etree
from docx import Document
from odf import text, table, teletype
from odf.opendocument import load

from .uniconstants import UNI_STRING

class UniExtractor:

	def extract(self, text):
		return re.findall(UNI_STRING, text)

	def add_extracted(self, strings, text):
		strings.extend(self.extract(text))

	def add_extracted_para(self, strings, para):
		self.add_extracted(strings, para.text)

	def add_extracted_table(self, strings, table):
		for row in table.rows:
			for cell in row.cells:
				for para in cell.paragraphs:
					self.add_extracted_para(strings, para)

	def extract_file(self, file):
		with open(file, 'r', encoding='utf-8') as f:
			text = f.read()
		return self.extract(text)

	def extract_html(self, file, classname=None):
		with open(file, 'r', encoding='utf-8') as f:
			doc = html.fromstring(f.read())
		strings = []
		if classname:
			xp = etree.XPath("//text()[ancestor::*[contains(concat(' ', normalize-space(@class), ' '), concat(' ', $cls, ' '))]]")
			for node in xp(doc, cls=classname):
				self.add_extracted(strings, str(node))
		else:
			for node in doc.xpath("//text()"):
				self.add_extracted(strings, str(node))
		return strings

	def extract_xml(self, file, attribute=None):
		doc = etree.parse(file)
		root = doc.getroot()
		strings = []
		if attribute:
			name, value = attribute
			xp = etree.XPath("//text()[ancestor::*[@name = $value]]")
			for node in xp(doc, name=name, value=value):
				self.add_extracted(strings, str(node))
		else:
			for node in root.xpath("//text()"):
				self.add_extracted(strings, str(node))
		return strings

	def extract_docx(self, file):
		doc = Document(file)
		strings = []
		for para in doc.paragraphs:
			self.add_extracted_para(strings, para)
		for table in doc.tables:
			self.add_extracted_table(strings, table)
		return strings
		
	def extract_odt(self, file):
		doc = load(file)
		strings = []
		for para in doc.body.getElementsByType(text.P):
			p_text = teletype.extractText(para)
			self.add_extracted(strings, p_text)
		return strings
